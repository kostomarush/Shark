from django.shortcuts import render, redirect, get_object_or_404
from .models import ScanInfo, DataServer, SegmentScan, IPAddress, SegmentResult, ResultPorts, CveInformation, ResultPortsAim, CveInformationAim, LevelCveAim, LevelCve
from .forms import DataServerForm, SegmentScanForm
from django.http import HttpResponse
from tempfile import NamedTemporaryFile
from django.contrib.auth.decorators import login_required
import docx
import ipaddress
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH

@login_required(redirect_field_name=None, login_url='/')
def tcp(request):
    return render(request, 'server/tcp.html')

@login_required(redirect_field_name=None, login_url='/')
def udp(request):
    return render(request, 'server/udp.html') 

@login_required(redirect_field_name=None, login_url='/')
def os(request):
    return render(request, 'server/os.html') 

def count_ports_for_report(item):
    # Get the SegmentResult object
    segment_result = SegmentResult.objects.get(id=item)

    # Count the number of ports for the given SegmentResult
    ports_count = ResultPorts.objects.filter(all_info=segment_result).count()

    return ports_count

def count_cves_for_segment_result(segment_result_id):
    
    segment_result = SegmentResult.objects.get(id=segment_result_id)

    # Count the number of ports for the given SegmentResult
    cve_count = LevelCve.objects.filter(result=segment_result).count()

    return cve_count

def find_ports(item_port):
    
    segment_result = SegmentResult.objects.get(id=item_port)
    
    all_data = ResultPorts.objects.filter(all_info=segment_result).all()

    return all_data

def find_cve(cve):
    
    segment_result = SegmentResult.objects.get(id=cve)
    
    all_data = ResultPorts.objects.filter(all_info=segment_result)
    
    info = []
    
    for result in all_data:
    
        cve_item = CveInformation.objects.filter(result_ports=result).all()
        
        for i in cve_item:
            info.append(i.cve_information)
    
    return info


def generate_word_report(request, pk):
    task = SegmentScan.objects.get(pk = pk)
    all_ip_addresses = IPAddress.objects.all()
    ports_by_host = {}
    segment_results_by_segment = {}

    for all_ip_address in all_ip_addresses:

        segment_results = SegmentResult.objects.filter(result=all_ip_address)


        segment_results_by_segment[all_ip_address] = segment_results

    items = []
    for all_ip_address, segment_results in segment_results_by_segment.items():
        if all_ip_address.seg_scan == task:
            for result in segment_results:
                items.append(result)


    for segment_result in segment_results:
        # Получите связанные с этим объектом IPAddress
        ports = ResultPorts.objects.filter(all_info=segment_result)

        # Сохраните их в словаре
        ports_by_host[segment_result] = ports

    table = []
    for ports, segment_result in ports_by_host.items():
        if ports == task:
            for ports in segment_result:
                table.append(ports)
    

    # Извлекаем данные из базы данных
    task = SegmentScan.objects.get(pk = pk)
    # Создаем новый документ Word
    document = docx.Document()
    # Добавляем раздел "header" в документ
    header_section = document.sections[0]
    header = header_section.header

    # Вставляем изображение в правый верхний угол
    image_path = '/home/user/Dipl/djnagoPRC/djangoRPC/server/static/server/images/logo-shark_2.png'  # Укажите путь к вашему изображению
    header.paragraphs[0].add_run().add_picture(image_path, width=Inches(1.5))  # Вы можете настроить ширину изображения
    title = document.add_heading(f'Отчет по безопасности сети {task.ip}/{task.mask}', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    styles = document.styles
    styles['Heading 1'].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    styles['Heading 2'].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    styles['Heading 3'].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    styles['Heading 4'].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    # Добавляем информацию о дате сканирования и версии сканера
    document.add_heading("Служебная информация:", level=2)
    scanner_version = '1.0'  # Замените на реальную версию вашего сканера
    document.add_paragraph(f"Дата сканирования: {task.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph(f"Версия сканера: {scanner_version}")
    document.add_paragraph(f"Режим сканирования: {task.mode}")
    document.add_paragraph(f"Режим CVE: {task.cve_report}")
    document.add_paragraph(f"Режим сканирования всех портов: {task.full_scan}")
        
    # Добавляем информацию о хостах в документ
    document.add_heading('Общая информация о сетевых узлах:', level=2)
    
    table = document.add_table(rows=1, cols=4, style='Table Grid')
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.cell(0, 0).text = 'Ip Адрес узла'
    table.cell(0, 1).text = 'Состояние хоста'
    table.cell(0, 2).text = 'Количество открытых портов'
    table.cell(0, 3).text = 'Общее количество CVE'
    
    for item in items:
        row_cells = table.add_row().cells
        row_cells[0].text = item.host
        row_cells[1].text = item.state_scan
        row_cells[2].text = str(count_ports_for_report(item.id))
        row_cells[3].text = str(count_cves_for_segment_result(item.id))


    document.add_heading('Подробная информация о сетевых узлах:', level=2)
    
    for host in items:
        document.add_heading(f'Узел {host.host}/{task.mask}', level=3)
        document.add_paragraph(f"Состояние: {host.state_scan}")
        document.add_paragraph(f"Количество открытых портов: {count_ports_for_report(host.id)}")
        document.add_paragraph(f"Общее количество CVE: {count_cves_for_segment_result(host.id)}")
        if task.mode == 'OS':
            document.add_paragraph(f"Операционная система: {host.full_name}")
            document.add_paragraph(f"Вендор: {host.vendor}")
            document.add_paragraph(f"Семейство: {host.osfamily}")
            document.add_paragraph(f"Версия: {host.osgen}")
            document.add_paragraph(f"Точность сканирования: {host.accuracy}")
        
        document.add_heading('Таблица информации о портах:', level=4)
        table = document.add_table(rows=1, cols=5, style='Table Grid')
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table.cell(0, 0).text = 'Номер порта'
        table.cell(0, 1).text = 'Состояние'
        table.cell(0, 2).text = 'Причина'
        table.cell(0, 3).text = 'Сервис'
        table.cell(0, 4).text = 'CVE'
        
        table.columns[4].width = Cm(14)
        
        for i in find_ports(host.id):
            row_cells = table.add_row().cells
            row_cells[0].text = i.port
            row_cells[1].text = i.state
            row_cells[2].text = i.reason
            row_cells[3].text = i.service
            row_cells[4].text = i.one_cve
            
        document.add_heading('Описание CVE:', level=4)
        paragraph = document.add_paragraph()
        for i in find_cve(host.id):
            run_mono2 = paragraph.add_run(f"{i}")
            font_mono2 = run_mono2.font
            font_mono2.name = 'Courier New'
            font_mono2.size = Pt(10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
    # Настраиваем стили форматирования
    
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(12)

    # Сохраняем документ во временном файле
    temp_file_path = "report.docx"
    document.save(temp_file_path)

    # Отправляем файл пользователю как вложение
    with open(temp_file_path, 'rb') as f:
        response = HttpResponse(f.read(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response['Content-Disposition'] = f'attachment; filename=report.docx'

    return response

@login_required(redirect_field_name=None, login_url='/')
def dashboard(request):
    return render(request, 'server/dashboard.html')

def not_found (request):
    return render(request, 'server/404.html')

@login_required(redirect_field_name=None, login_url='/')
def detail_seg(request, pk):
    item = get_object_or_404(SegmentScan, pk=pk)
    segment_scans = SegmentScan.objects.all()
    all_ip_addresses = IPAddress.objects.all()
    task_done = IPAddress.objects.filter(seg_scan = item, tag='Done').count()
    tasks = IPAddress.objects.filter(seg_scan = item).count()
    # Создайте словарь, где ключами будут объекты SegmentScan, а значениями будут связанные IP-адреса
    ip_addresses_by_segment = {}
    segment_results_by_segment = {}

    for all_ip_address in all_ip_addresses:

        segment_results = SegmentResult.objects.filter(result=all_ip_address)

        segment_results_by_segment[all_ip_address] = segment_results

    result_dict = []
    for all_ip_address, segment_results in segment_results_by_segment.items():
        if all_ip_address.seg_scan == item:
            for result in segment_results:
                result_dict.append(result)

    # Пройдите по каждому объекту SegmentScan
    for segment_scan in segment_scans:
        # Получите связанные с этим объектом IPAddress
        ip_addresses = IPAddress.objects.filter(seg_scan=segment_scan)
        # Сохраните их в словаре
        ip_addresses_by_segment[segment_scan] = ip_addresses
    ip_dict = []
    for segment_scan, ip_addresses in ip_addresses_by_segment.items():
        if segment_scan == item:
            for ip_address in ip_addresses:
                ip_dict.append(ip_address)
                
    vulnerability_counts_by_year = {}

    unique_years = LevelCve.objects.filter(result__result__seg_scan=item).values('year').distinct()
    # Итерируемся по уникальным годам и подсчитываем уровни уязвимостей
    for year_info in unique_years:
        year = year_info['year']


        critical_count = LevelCve.objects.filter(year=year, level='Критичная').count()


        high_count = LevelCve.objects.filter(year=year, level='Высокая').count()


        medium_count = LevelCve.objects.filter(year=year, level='Средняя').count()
        
        
        normal_count = LevelCve.objects.filter(year=year, level='Низкая').count()

        # Создаем словарь для уровней уязвимостей текущего года
        levels_dict = {
            'Критичная': critical_count,
            'Высокая': high_count,
            'Средняя': medium_count,
            'Низкая': normal_count
        }
        vulnerability_counts_by_year[year] = levels_dict

    client_1 = 0
    client_3 = 0
    client_2 = 0
    data_seg = IPAddress.objects.filter(seg_scan=item).in_bulk()
    for id in data_seg:
            if data_seg[id].tag == 'Done' and data_seg[id].client.ip_client == '1':
                client_1 += 1
            if data_seg[id].tag == 'Done' and data_seg[id].client.ip_client == '2':
                client_2 += 1
            if data_seg[id].tag == 'Done' and data_seg[id].client.ip_client == '3':
                client_3 += 1
                
    return render(request, 'server/detail_seg.html', {'item': item,
                                                      'all_ip': ip_dict,
                                                      'result': result_dict,
                                                      'task_done': task_done,
                                                      'cve_year': vulnerability_counts_by_year,
                                                      'client_1': client_1,
                                                      'client_2': client_2,
                                                      'client_3': client_3
                                                      })


@login_required(redirect_field_name=None, login_url='/')
def remove_item(request, pk):
    item = DataServer.objects.get(pk=pk)
    item.delete()
    return redirect('aim')


@login_required(redirect_field_name=None, login_url='/')
def remove_segment(request, pk):
    item = SegmentScan.objects.get(pk=pk)
    item.delete()
    return redirect('segment')


@login_required(redirect_field_name=None, login_url='/')
def data(request):
    query_results = ScanInfo.objects.all()
    data_serv = DataServer.objects.all()
    task_done = DataServer.objects.filter(tag='Done').count()
    error = ''
    if request.method == 'POST':
        form = DataServerForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('aim')
        else:
            error = 'Форма не верна'
    form = DataServerForm()
    tasks = {
        'form': form,
        'data_serv': data_serv,
        'error': error,
        'section': query_results,
        'task_done': task_done,
    }
    return render(request, 'server/aim.html', tasks)


@login_required(redirect_field_name=None, login_url='/')
def segment(request):
    
    scan_segment = SegmentScan.objects.all()
    
    error = ''
    if request.method == 'POST':
        form = SegmentScanForm(request.POST)
        if form.is_valid():
            segment_scan_instance = form.save()
            net = segment_scan_instance.ip
            mask = segment_scan_instance.mask
            network = ipaddress.IPv4Network(f'{net}/{mask}', strict=False)
            segments = [
                ipaddr for ipaddr in network.subnets(prefixlen_diff=3)]
            for addr in segments:
                IPAddress.objects.create(
                    address=f'{addr}', seg_scan=segment_scan_instance)
            
        else:
            error = 'Форма не верна'

    form_segment = SegmentScanForm()
    if scan_segment:
        for scan_res in scan_segment:
        
                all_done = IPAddress.objects.filter(seg_scan = scan_res,
                    tag='Done').count() == IPAddress.objects.filter(seg_scan = scan_res).count()
    
                if all_done:
                
                    segment_scan = SegmentScan.objects.get(id=scan_res.id)
                    segment_scan.state_scan = 'Done'
                    segment_scan.save()

    seg = {
        'form_segment': form_segment,
        'error': error,
        'scan_segment': scan_segment,
    }
    return render(request, 'server/segment.html', seg)


def port_information(request, pk):
    item = get_object_or_404(SegmentResult, pk=pk)
    segment_results = SegmentResult.objects.all()
    ports_by_host = {}
    
    cve_level_host = {}
    
    for cve_level in segment_results:

        level = LevelCve.objects.filter(result=cve_level)

        # Сохраните их в словаре
        cve_level_host[cve_level] = level

    level_dict = []
    for level, cve_level in cve_level_host.items():
        if level == item:
            num_rows = LevelCve.objects.filter(result = level).count()
            critical = LevelCve.objects.filter(result = level,level='Критичная').count()
            high = LevelCve.objects.filter(result = level,level='Высокая').count()
            medium = LevelCve.objects.filter(result = level,level='Средняя').count()
            normal = LevelCve.objects.filter(result = level,level='Низкая').count()
            for level in cve_level:
                level_dict.append(level)
    # Пройдите по каждому объекту SegmentScan
    for segment_result in segment_results:
        # Получите связанные с этим объектом IPAddress
        ports = ResultPorts.objects.filter(all_info=segment_result)

        # Сохраните их в словаре
        ports_by_host[segment_result] = ports

    port_dict = []
    for ports, segment_result in ports_by_host.items():
        if ports == item:
            open = ResultPorts.objects.filter(all_info = ports, state='open').count()
            filtered = ResultPorts.objects.filter(all_info = ports, state='filtered').count()
            close = ResultPorts.objects.filter(all_info = ports, state='closed').count()
            open_filtered = ResultPorts.objects.filter(all_info = ports, state='open|filtered').count()
            for ports in segment_result:
                port_dict.append(ports)

    return render(request, 'server/port_information.html', {'item': item,
                                                         'port_dict': port_dict,
                                                         'level_dict': level_dict,
                                                         'count': num_rows,
                                                        'critical': critical,
                                                        'high': high,
                                                        'medium': medium,
                                                        'normal': normal,
                                                        'open':open,
                                                        'filtered': filtered,
                                                        'close':close,
                                                        'open_filtered':open_filtered})



def cve_information(request, pk):
    item = get_object_or_404(ResultPorts, pk=pk)
    port_results = ResultPorts.objects.all()
    cve_by_port = {}
    # Пройдите по каждому объекту SegmentScan
    for port_result in port_results:
        # Получите связанные с этим объектом IPAddress
        cve = CveInformation.objects.filter(result_ports=port_result)

        # Сохраните их в словаре
        cve_by_port[port_result] = cve

    cve_dict = []
    for cve, port_result in cve_by_port.items():
        if cve == item:
            for i in port_result:
                cve_dict.append(i)

    return render(request, 'server/cve_information.html', {'item': item,
                                                            'cve_dict': cve_dict})



def port_info_aim(request, pk):
    item = get_object_or_404(ScanInfo, pk=pk)
    scan_results = ScanInfo.objects.all()
    
    cve_level_host = {}
    ports_by_host = {}
    
    
    for cve_level in scan_results:

        level = LevelCveAim.objects.filter(result=cve_level)

        # Сохраните их в словаре
        cve_level_host[cve_level] = level

    level_dict = []
    for level, cve_level in cve_level_host.items():
        if level == item:
            num_rows = LevelCveAim.objects.filter(result = level).count()
            critical = LevelCveAim.objects.filter(result = level,level='Критичная').count()
            high = LevelCveAim.objects.filter(result = level,level='Высокая').count()
            medium = LevelCveAim.objects.filter(result = level,level='Средняя').count()
            normal = LevelCveAim.objects.filter(result = level,level='Низкая').count()
            for level in cve_level:
                level_dict.append(level)
    
    
    # Пройдите по каждому объекту SegmentScan
    for scan_result in scan_results:
        # Получите связанные с этим объектом IPAddress
        ports = ResultPortsAim.objects.filter(all_info=scan_result)

        # Сохраните их в словаре
        ports_by_host[scan_result] = ports

    port_dict = []
    for ports, scan_result in ports_by_host.items():
        if ports == item:
            open = ResultPortsAim.objects.filter(all_info = ports, state='open').count()
            filtered = ResultPortsAim.objects.filter(all_info = ports, state='filtered').count()
            close = ResultPortsAim.objects.filter(all_info = ports, state='closed').count()
            open_filtered = ResultPortsAim.objects.filter(all_info = ports, state='open|filtered').count()
            for ports in scan_result:
                port_dict.append(ports)

    return render(request, 'server/port_info_aim.html', {'item': item,
                                                         'port_dict': port_dict,
                                                         'level_dict': level_dict,
                                                         'count': num_rows,
                                                        'critical': critical,
                                                        'high': high,
                                                        'medium': medium,
                                                        'normal': normal,
                                                        'open':open,
                                                        'filtered': filtered,
                                                        'close':close,
                                                        'open_filtered':open_filtered})

def cve_information_aim(request, pk):
    item = get_object_or_404(ResultPortsAim, pk=pk)
    port_results = ResultPortsAim.objects.all()

    cve_by_port = {}
    # Пройдите по каждому объекту SegmentScan
    for port_result in port_results:
        # Получите связанные с этим объектом IPAddress
        cve = CveInformationAim.objects.filter(result_ports=port_result)

        # Сохраните их в словаре
        cve_by_port[port_result] = cve

    cve_dict = []
    for cve, port_result in cve_by_port.items():
        if cve == item:
            for i in port_result:
                cve_dict.append(i)

    return render(request, 'server/cve_information_aim.html', {'item': item,
                                                            'cve_dict': cve_dict})